"""
ClauseIQ — Text Extraction Module

Extracts raw text from PDF (PyMuPDF) and DOCX (python-docx) files
with page-level boundaries preserved per AGENT.md Section 9.2.
"""

from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.core.logging_config import get_logger

logger = get_logger("extractor")


@dataclass
class PageContent:
    """Represents extracted text from a single page."""
    page_number: int
    text: str


@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""
    pages: list[PageContent] = field(default_factory=list)
    page_count: int = 0
    full_text: str = ""


def extract_from_pdf(file_path: str) -> ExtractionResult:
    """
    Extract text from a PDF file using PyMuPDF.

    Preserves page-level boundaries for citation support.
    """
    doc = fitz.open(file_path)
    pages: list[PageContent] = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text("text")
        if text.strip():
            pages.append(PageContent(page_number=page_num + 1, text=text))

    doc.close()

    full_text = "\n\n".join(p.text for p in pages)
    result = ExtractionResult(
        pages=pages,
        page_count=len(doc) if hasattr(doc, '__len__') else len(pages),
        full_text=full_text,
    )

    logger.info(
        "PDF extracted: path=%s, pages=%d, chars=%d",
        file_path, len(pages), len(full_text),
    )
    return result


def extract_from_docx(file_path: str) -> ExtractionResult:
    """
    Extract text from a DOCX file using python-docx.

    DOCX files don't have native page boundaries, so we treat
    the entire document as page 1 and split by paragraphs.
    """
    doc = DocxDocument(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)

    # DOCX doesn't have page numbers — we assign page 1 to everything
    pages = [PageContent(page_number=1, text=full_text)] if full_text else []

    result = ExtractionResult(
        pages=pages,
        page_count=1,
        full_text=full_text,
    )

    logger.info(
        "DOCX extracted: path=%s, paragraphs=%d, chars=%d",
        file_path, len(paragraphs), len(full_text),
    )
    return result


def extract_text(file_path: str, file_type: str) -> ExtractionResult:
    """
    Dispatch extraction based on file type.

    Args:
        file_path: Path to the uploaded file.
        file_type: Either "pdf" or "docx".

    Returns:
        ExtractionResult with page-level text.

    Raises:
        ValueError: If file type is unsupported.
    """
    if file_type == "pdf":
        return extract_from_pdf(file_path)
    elif file_type == "docx":
        return extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
