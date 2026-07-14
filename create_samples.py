import docx

doc = docx.Document()
doc.add_heading('Sample Software License Agreement', 0)

doc.add_heading('1. Grant of License', level=1)
doc.add_paragraph('Licensor hereby grants to Licensee a non-exclusive, non-transferable license to use the Software for internal business purposes.')

doc.add_heading('2. Term and Termination', level=1)
doc.add_paragraph('This Agreement shall commence on the Effective Date and remain in effect until terminated. Either party may terminate this Agreement upon 30 days written notice.')

doc.add_heading('3. Limitation of Liability', level=1)
doc.add_paragraph('In no event shall Licensor be liable for any indirect, incidental, or consequential damages arising out of the use of the Software.')

doc.add_heading('4. Governing Law', level=1)
doc.add_paragraph('This Agreement shall be governed by the laws of the State of California.')

doc.save('C:/Users/Administrator/Desktop/clauseiq/sample_contract_a.docx')

# Create a second contract for comparison
doc2 = docx.Document()
doc2.add_heading('Sample Software License Agreement - Modified', 0)

doc2.add_heading('1. Grant of License', level=1)
doc2.add_paragraph('Licensor hereby grants to Licensee an exclusive, transferable license to use the Software for any purpose.')

doc2.add_heading('2. Term and Termination', level=1)
doc2.add_paragraph('This Agreement shall commence on the Effective Date and remain in effect for 12 months. Either party may terminate this Agreement upon 60 days written notice.')

doc2.add_heading('3. Confidentiality', level=1)
doc2.add_paragraph('Both parties agree to maintain strict confidentiality regarding all proprietary information shared during the term of this Agreement.')

doc2.save('C:/Users/Administrator/Desktop/clauseiq/sample_contract_b.docx')

print("Created sample_contract_a.docx and sample_contract_b.docx")
